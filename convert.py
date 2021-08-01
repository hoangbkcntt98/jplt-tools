import textract
from read_res import readAnswers
from os import walk
from tkinter import Tk, Frame
import docx
answers = readAnswers()


def collectQuizPaths():
    filenames = next(walk('./debai'), (None, None, []))[2]  # [] if no file
    path = ['./debai/'+i for i in filenames]
    return path


def readFile(path):
    f = open(path, 'rb')
    document = Document(f)
    f.close()
    return document


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def writeFile(filename, text):
    file = open(filename, 'ab')
    file.write(text.encode('utf-8'))


def saveAsDocx(filename, text):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)


def trimString(string):
    trimString = string.split(" ")
    trimString = "".join(trimString)
    trimString = trimString.strip()
    return trimString


def findNextMondai(mondai):
    afterMondai = mondai.split(mondai[-1])
    numAfter = int(mondai[-1])+1
    return afterMondai[0]+str(numAfter)


def getExcercise(excercise, text):
    textStr = trimString(text)
    afterExcer = str(int(excercise)+1)
    mondai = "問題"+excercise
    afterMondai = "問題"+afterExcer
    if textStr.find(afterMondai) != -1:
        textStr = textStr.split(afterMondai)
    textStr = trimString(textStr[0])
    # textStr = textStr.decode('utf-8')
    textStr = textStr.split(mondai)

    # print(textStr)
    textStr = mondai+textStr[1]
    textStr = trimString(textStr)
    # print(textStr)
    lines = textStr.splitlines()
    excerciseData = []
    for line in lines:
        # print(line)
        if len(line) > 0:
            excerciseData.append(line)
    excerciseData = excerciseData[1:]
    return excerciseData


def getMondai3(mondai, text, mondainext=None):
    textStr = trimString(text)
    if mondainext == None:
        afterMondai = findNextMondai(mondai)
    else:
        afterMondai = mondainext
    # print(afterMondai)
    # print(mondai.encode('utf-8'))
    # print(textStr)
    if textStr.find(afterMondai) != -1:
        # print('no')
        textStr = textStr.split(afterMondai)
        textStr = trimString(textStr[0])
    # textStr = textStr.decode('utf-8')
    textStr = textStr.split(mondai)

    # print(textStr)
    textStr = mondai+textStr[1]
    textStr = trimString(textStr)
    # print(textStr)
    lines = textStr.splitlines()
    excerciseData = []
    for line in lines:
        # print(line)
        if len(line) > 0:
            excerciseData.append(line)
    excerciseData = excerciseData[1:]
    return excerciseData


def openFile(path):
    text = textract.process(path)
    textStr = text.decode('utf-8')
    return textStr


def readData(path, excercise):
    # Open excercise files
    # textStr = openFile(path)
    jp_number = ["１", "２", "３", "４", "５", "６", "７", "８", "９", "１０", "１１", "１２"]
    dict_num = {}
    for count, num in enumerate(jp_number, start=1):
        dict_num[str(num)] = str(count)

    print(dict_num)
    textStr = getText(path)
    [textStr.replace(t, dict_num[t])
     for t in jp_number if textStr.find(t) != -1]
    time = path.split('.docx')[0].split('/')[-1]
    temp = answers[time]
    mondai3 = []
    try:
        if excercise == None:
        mondai3 = getMondai3('問題3', textStr)
        else:

    except:
        mondai3 = getMondai3('問題３', textStr, "問題４")

    # print(temp)
    # Map excercise with result
    resultsInExer = []
    excercises = []
    for count, item in enumerate(mondai3, start=1):
        if count % 2 == 0:
            num = int(count/2-1)
            # print(temp[num])
            # print(item)
            index = item.find(str(temp[num]))
            # print(index)
            resultsInExer.append(item[index+1])
        else:
            excercises.append(item)
    finalRes = []
    for count, exc in enumerate(excercises, start=0):
        # print(resultsInExer[count])
        # print(exc.find("()"))
        if exc.find("()") != -1:
            exc = exc.replace("()", resultsInExer[count])
        if exc.find("(。)") != -1:
            exc = exc.replace("(。)", resultsInExer[count])
        if exc.find("（）") != -1:
            exc = exc.replace("（）", resultsInExer[count])
        if exc.find("（　）") != -1:
            exc = exc.replace("（　）", resultsInExer[count])

        finalRes.append(exc)
    # print(finalRes)
    finalRes.append(time)
    return finalRes


def readMultileFile(list, excercise):
    data = []
    for path in list:
        if excercise != None:
            data.append(readData(path, excercise))
        else:
            data.append(readData(path, None))
    return data


def findWord(self, word, excercise=None):

    res = []
    paths = collectQuizPaths()
    data = readMultileFile(paths, excercise)

    for d in data:
        time = d.pop()
        for y in d:
            if y.find(word) != -1 and y != time:
                res.append([y, time])
    # print(res)
    return res


def findWord4(self, word):
    res = []
    paths = collectQuizPaths()
    data = readMultileFile(paths, 4)


findWord(None, '')
